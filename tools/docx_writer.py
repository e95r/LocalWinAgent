"""Помощники для работы с документами Word (.docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _add_paragraphs(document: Document, content: str, *, replace_first: bool) -> None:
    lines = content.splitlines() if content else [""]
    for index, line in enumerate(lines):
        if replace_first and index == 0 and document.paragraphs:
            document.paragraphs[0].text = line
        else:
            document.add_paragraph(line)


def write_docx(path: str | Path, content: str) -> None:
    """Создать или перезаписать документ .docx указанным содержимым."""

    target = Path(path)
    _ensure_parent(target)
    document = Document()
    _add_paragraphs(document, content, replace_first=True)
    document.save(target)


def append_docx(path: str | Path, content: str) -> None:
    """Добавить текст в существующий документ .docx, создавая при необходимости."""

    target = Path(path)
    _ensure_parent(target)
    if target.exists():
        document = Document(target)
        replace_first = False
    else:
        document = Document()
        replace_first = True
    _add_paragraphs(document, content, replace_first=replace_first)
    document.save(target)
